"""Generate a folder of sample resumes for testing the screener.

Creates a realistic mix under tests/sample_resumes/:
  - Text-based PDF resumes with varying relevance to a backend/Python/AWS role
  - A DOCX resume that puts skills in a TABLE (to test table extraction)
  - A "scanned" image-only PDF with NO extractable text (to test failure reporting)
  - A legacy .doc file (to test that .doc is reported, not silently dropped)

Run with:  uv run python tests/generate_samples.py
"""

from __future__ import annotations

import os

from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

HERE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(HERE, "sample_resumes")


def write_pdf(filename: str, lines: list[str]) -> None:
    path = os.path.join(OUT, filename)
    c = canvas.Canvas(path, pagesize=letter)
    width, height = letter
    y = height - 60
    for line in lines:
        if y < 60:
            c.showPage()
            y = height - 60
        c.setFont("Helvetica", 11)
        c.drawString(60, y, line[:110])
        y -= 16
    c.save()


def write_scanned_pdf(filename: str) -> None:
    """A PDF with only a drawn rectangle — zero extractable text.

    Simulates a scanned/image resume. pypdf will extract ~nothing, which must
    trip the <100-char quality gate and be reported as unreadable."""
    path = os.path.join(OUT, filename)
    c = canvas.Canvas(path, pagesize=letter)
    width, height = letter
    c.rect(60, 60, width - 120, height - 120, stroke=1, fill=0)
    c.line(60, height - 200, width - 60, height - 200)
    c.save()


def write_docx_with_table(filename: str) -> None:
    path = os.path.join(OUT, filename)
    doc = Document()
    doc.add_heading("Priya Sharma", level=1)
    doc.add_paragraph("Senior Backend Engineer | Bangalore, India")
    doc.add_paragraph(
        "Summary: Backend engineer with 7 years building scalable Python "
        "services on AWS. Led migration of a monolith to microservices."
    )
    doc.add_paragraph("Skills (see table):")
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Category"
    hdr[1].text = "Technologies"
    rows = [
        ("Languages", "Python, Go, SQL"),
        ("Cloud", "AWS (EC2, Lambda, S3, RDS, ECS), Terraform"),
        ("Frameworks", "FastAPI, Django, Flask"),
        ("Data", "PostgreSQL, Redis, Kafka"),
        ("Practices", "CI/CD, Docker, Kubernetes, observability"),
    ]
    for cat, tech in rows:
        cells = table.add_row().cells
        cells[0].text = cat
        cells[1].text = tech
    doc.add_paragraph("")
    doc.add_paragraph(
        "Experience: Tech Lead at FintechCo (2019-present). Designed event-driven "
        "payment pipeline handling 5M transactions/day. Mentored 4 engineers."
    )
    doc.add_paragraph(
        "Education: B.Tech Computer Science, IIT Bombay, 2016."
    )
    doc.save(path)


def write_legacy_doc(filename: str) -> None:
    """A file with a .doc extension. Content is irrelevant — the screener
    should detect the extension and report it as legacy/unsupported."""
    path = os.path.join(OUT, filename)
    with open(path, "wb") as f:
        f.write(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1")  # OLE2 magic bytes
        f.write(b"This is a legacy binary .doc file. Not a real docx.\n" * 5)


RESUMES = {
    "alice_backend.pdf": [
        "Alice Johnson",
        "Senior Backend Engineer",
        "alice.johnson@example.com | San Francisco, CA",
        "",
        "SUMMARY",
        "Backend engineer with 8 years of experience building high-throughput",
        "Python services on AWS. Deep expertise in distributed systems, API design,",
        "and cloud infrastructure.",
        "",
        "SKILLS",
        "Python, FastAPI, Django, PostgreSQL, Redis, Kafka, Docker, Kubernetes,",
        "AWS (Lambda, ECS, S3, RDS, DynamoDB), Terraform, CI/CD, microservices.",
        "",
        "EXPERIENCE",
        "Staff Engineer, CloudScale Inc (2020-present)",
        "- Architected a microservices platform serving 10M requests/day on AWS ECS.",
        "- Reduced p99 latency 40% by introducing Redis caching and async workers.",
        "Backend Engineer, DataWorks (2016-2020)",
        "- Built ETL pipelines in Python processing terabytes of event data.",
        "",
        "EDUCATION",
        "B.S. Computer Science, UC Berkeley, 2016.",
    ],
    "bob_fullstack.pdf": [
        "Bob Martinez",
        "Full-Stack Developer",
        "bob.martinez@example.com | Austin, TX",
        "",
        "SUMMARY",
        "Full-stack developer with 5 years building web apps. Comfortable across",
        "the stack with a frontend lean.",
        "",
        "SKILLS",
        "JavaScript, TypeScript, React, Node.js, Express, some Python, MongoDB,",
        "PostgreSQL, AWS basics (EC2, S3), Docker.",
        "",
        "EXPERIENCE",
        "Full-Stack Developer, ShopFast (2019-present)",
        "- Built React storefront and Node.js APIs for an e-commerce platform.",
        "- Some backend work in Python/Flask for internal tooling.",
        "",
        "EDUCATION",
        "B.S. Information Systems, UT Austin, 2018.",
    ],
    "carol_data.pdf": [
        "Carol Nguyen",
        "Data Engineer",
        "carol.nguyen@example.com | Seattle, WA",
        "",
        "SUMMARY",
        "Data engineer with 6 years building data platforms. Strong Python and SQL,",
        "heavy AWS data stack experience.",
        "",
        "SKILLS",
        "Python, SQL, Spark, Airflow, AWS (Glue, EMR, Redshift, S3, Lambda),",
        "Kafka, dbt, Snowflake, Terraform.",
        "",
        "EXPERIENCE",
        "Senior Data Engineer, StreamCo (2018-present)",
        "- Built real-time streaming pipelines on AWS handling billions of events.",
        "- Owned the data warehouse and orchestration with Airflow.",
        "",
        "EDUCATION",
        "M.S. Data Science, University of Washington, 2017.",
    ],
    "dave_frontend.pdf": [
        "Dave Wilson",
        "Frontend Engineer",
        "dave.wilson@example.com | New York, NY",
        "",
        "SUMMARY",
        "Frontend engineer specializing in design systems and accessible UI.",
        "",
        "SKILLS",
        "JavaScript, TypeScript, React, Vue, CSS, Figma, Webpack, Jest,",
        "design systems, accessibility (WCAG).",
        "",
        "EXPERIENCE",
        "Senior Frontend Engineer, DesignHub (2019-present)",
        "- Led the company design system used across 12 product teams.",
        "- Improved Lighthouse scores and accessibility compliance.",
        "",
        "EDUCATION",
        "B.A. Design, Parsons School of Design, 2015.",
    ],
    "emma_devops.pdf": [
        "Emma Brown",
        "DevOps / Platform Engineer",
        "emma.brown@example.com | Denver, CO",
        "",
        "SUMMARY",
        "Platform engineer with 7 years in cloud infrastructure and automation.",
        "Strong AWS and Kubernetes, scripts heavily in Python.",
        "",
        "SKILLS",
        "AWS (EKS, EC2, Lambda, CloudFormation), Kubernetes, Terraform, Python,",
        "Bash, Docker, Prometheus, Grafana, CI/CD, GitOps.",
        "",
        "EXPERIENCE",
        "Platform Engineer, InfraCloud (2018-present)",
        "- Managed multi-region Kubernetes clusters on AWS EKS.",
        "- Built internal developer platform and CI/CD pipelines in Python.",
        "",
        "EDUCATION",
        "B.S. Computer Engineering, Colorado State, 2016.",
    ],
    "frank_junior.pdf": [
        "Frank Lee",
        "Junior Software Engineer",
        "frank.lee@example.com | Remote",
        "",
        "SUMMARY",
        "Recent bootcamp grad eager to grow as a backend developer. 1 year",
        "of internship experience.",
        "",
        "SKILLS",
        "Python, Flask, SQL, Git, basic AWS (S3, EC2), HTML, CSS.",
        "",
        "EXPERIENCE",
        "Software Engineering Intern, StartupX (2023-2024)",
        "- Built REST endpoints in Flask and wrote unit tests.",
        "",
        "EDUCATION",
        "Coding Bootcamp, General Assembly, 2023.",
    ],
    "grace_nurse.pdf": [
        "Grace Kim",
        "Registered Nurse",
        "grace.kim@example.com | Chicago, IL",
        "",
        "SUMMARY",
        "Compassionate RN with 9 years in critical care. No software background.",
        "",
        "SKILLS",
        "Patient care, ICU, IV therapy, EHR (Epic), triage, team coordination.",
        "",
        "EXPERIENCE",
        "ICU Nurse, Mercy Hospital (2015-present)",
        "- Managed critical care for post-op patients.",
        "",
        "EDUCATION",
        "B.S. Nursing, University of Illinois Chicago, 2014.",
    ],
    "henry_principal.pdf": [
        "Henry Adams",
        "Principal Software Engineer",
        "henry.adams@example.com | Boston, MA",
        "",
        "SUMMARY",
        "Principal engineer with 12 years building cloud-native backend systems.",
        "Python and Go expert, extensive AWS architecture experience.",
        "",
        "SKILLS",
        "Python, Go, gRPC, FastAPI, PostgreSQL, Kafka, Redis, AWS (Lambda, ECS,",
        "EKS, RDS, DynamoDB, SQS, SNS), Terraform, Kubernetes, system design.",
        "",
        "EXPERIENCE",
        "Principal Engineer, ScaleNow (2017-present)",
        "- Led architecture for a platform processing 50M events/day on AWS.",
        "- Designed multi-tenant microservices and mentored 10+ engineers.",
        "Senior Engineer, BigTech (2012-2017)",
        "- Built backend services in Python and Go.",
        "",
        "EDUCATION",
        "M.S. Computer Science, MIT, 2012.",
    ],
}


def main() -> None:
    os.makedirs(OUT, exist_ok=True)
    for fname, lines in RESUMES.items():
        write_pdf(fname, lines)
    write_docx_with_table("priya_backend.docx")
    write_scanned_pdf("scanned_resume.pdf")
    write_legacy_doc("old_resume.doc")
    print(f"Wrote {len(RESUMES)} PDFs + 1 DOCX(table) + 1 scanned PDF + 1 .doc to {OUT}")


if __name__ == "__main__":
    main()
