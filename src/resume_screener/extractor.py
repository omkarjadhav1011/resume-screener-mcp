"""Pure text-extraction logic for resumes.

No MCP code lives here on purpose: these are plain functions so they stay
unit-testable and reusable by every tool in server.py.

Design rules (see APPENDIX B of the build plan):
  - These functions NEVER raise to the caller. Every failure mode becomes a
    structured value: extract_* return (text, error); extract_resume returns a
    record with ok=False and a human-readable `error`.
  - Silent data loss is the cardinal sin. A file we can't read must surface a
    reason, never just vanish.
"""

from __future__ import annotations

import os
import re
import unicodedata

from pypdf import PdfReader

from resume_screener.fields import parse_fields

# Below this many characters, normalized text is treated as "no real text"
# (almost always a scanned/image PDF or an effectively empty file). OCR is out
# of scope, so we report it rather than pretending we read a resume.
MIN_TEXT_CHARS = 100


def extract_pdf(path: str) -> tuple[str, str | None]:
    """Extract text from a PDF. Returns (text, error).

    Handles encrypted PDFs (tries an empty-password decrypt) and never raises:
    any failure becomes ("", reason)."""
    try:
        reader = PdfReader(path)
        if reader.is_encrypted:
            try:
                # Many "encrypted" PDFs use an empty owner password and open fine.
                result = reader.decrypt("")
                if result == 0:  # PasswordType.NOT_DECRYPTED
                    return "", "password-protected PDF"
            except Exception:
                return "", "password-protected PDF"
        parts: list[str] = []
        for page in reader.pages:
            try:
                parts.append(page.extract_text() or "")
            except Exception:
                # One bad page shouldn't kill the whole document.
                continue
        return "\n".join(parts), None
    except Exception as exc:
        return "", f"PDF read error: {type(exc).__name__}"


def extract_docx(path: str) -> tuple[str, str | None]:
    """Extract text from a .docx. Returns (text, error).

    Pulls BOTH paragraphs AND table cells — resumes frequently put skills and
    dates in tables, and missing them loses half the content. A legacy binary
    .doc mislabeled/handed here will fail to open; we report that clearly."""
    try:
        # Imported lazily so a docx import problem can't break PDF-only runs.
        from docx import Document

        doc = Document(path)
        parts: list[str] = [p.text for p in doc.paragraphs if p.text]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        parts.append(cell.text)
        return "\n".join(parts), None
    except Exception as exc:
        # python-docx raises PackageNotFoundError for legacy .doc / non-zip files.
        name = type(exc).__name__
        if name == "PackageNotFoundError":
            return (
                "",
                "legacy .doc format not supported — convert to .docx or PDF",
            )
        return "", f"DOCX read error: {name}"


_CONTROL_CHARS = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")
_WS_RUN = re.compile(r"[ \t]+")
_NEWLINE_RUN = re.compile(r"\n\s*\n\s*")


def normalize_text(raw: str) -> str:
    """Deterministic cleanup. No 'smart' rewriting — just consistent text.

    NFKC-normalizes unicode (accents/ligatures), strips control chars, and
    collapses whitespace runs so downstream TF-IDF and Claude see clean text."""
    if not raw:
        return ""
    text = unicodedata.normalize("NFKC", raw)
    text = _CONTROL_CHARS.sub("", text)
    # Normalize line endings, then collapse spaces/tabs and blank-line runs.
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = _WS_RUN.sub(" ", text)
    text = _NEWLINE_RUN.sub("\n", text)
    # Trim trailing spaces on each line.
    text = "\n".join(line.strip() for line in text.split("\n"))
    return text.strip()


def extract_resume(path: str) -> dict:
    """The single entry point every tool uses. NEVER raises.

    Dispatches on extension, normalizes, applies the quality gate, and returns a
    stable record. Every failure mode becomes a record with ok=False and a
    human-readable `error`."""
    filename = os.path.basename(path)
    record = {
        "filename": filename,
        "path": os.path.abspath(path),
        "text": "",
        "char_count": 0,
        "ok": False,
        "error": None,
    }

    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        text, error = extract_pdf(path)
    elif ext == ".docx":
        text, error = extract_docx(path)
    elif ext == ".doc":
        record["error"] = (
            "legacy .doc format not supported — convert to .docx or PDF"
        )
        return record
    else:
        record["error"] = f"unsupported file type: {ext or '(none)'}"
        return record

    if error:
        record["error"] = error
        return record

    clean = normalize_text(text)
    if len(clean) < MIN_TEXT_CHARS:
        record["error"] = (
            "no extractable text (likely scanned/image PDF — OCR not supported)"
        )
        return record

    record["text"] = clean
    record["char_count"] = len(clean)
    record["ok"] = True
    # Additive: attach cheap structured fields (name/email/phone/links/years).
    # Skills stay empty here — they're computed on demand against a JD-derived
    # vocabulary by later tools. Never breaks existing record consumers.
    record["fields"] = parse_fields(clean)
    return record


# Extensions we recognize as resumes for inventory purposes.
_RESUME_EXTS = {".pdf", ".docx"}


def find_resume_files(folder: str) -> tuple[list[str], list[str], str | None]:
    """Find resume files in a folder.

    Returns (resume_paths, legacy_doc_paths, error). `resume_paths` are .pdf/.docx
    sorted deterministically; `legacy_doc_paths` are .doc files reported separately
    so they don't silently vanish. `error` is set (and lists empty) when the path
    is missing or not a directory."""
    if not os.path.exists(folder):
        return [], [], f"That folder doesn't exist: {folder}"
    if not os.path.isdir(folder):
        return [], [], f"That path isn't a folder: {folder}"

    resumes: list[str] = []
    legacy: list[str] = []
    try:
        entries = os.listdir(folder)
    except Exception as exc:
        return [], [], f"Couldn't read that folder ({type(exc).__name__}): {folder}"

    for name in entries:
        full = os.path.join(folder, name)
        if not os.path.isfile(full):
            continue
        ext = os.path.splitext(name)[1].lower()
        if ext in _RESUME_EXTS:
            resumes.append(full)
        elif ext == ".doc":
            legacy.append(full)

    resumes.sort()
    legacy.sort()
    return resumes, legacy, None
